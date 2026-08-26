import os
import time
import uuid
import queue
import shutil
import sqlite3
import re
import subprocess
import threading
from datetime import datetime
from docx import Document

from fastapi import FastAPI, UploadFile, File, Form, Request, Response, HTTPException
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, RedirectResponse, PlainTextResponse
import uvicorn
import telebot
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton

# --- КОНФИГУРАЦИЯ ---
TG_BOT_TOKEN = os.environ.get("TG_BOT_TOKEN", "8504609196:AAE-AXIpfytvvDigddCHMvTT9ukPp9m-SWw")
TG_BOT_USERNAME = os.environ.get("TG_BOT_USERNAME", "whisper_log_bot")
SITE_URL = os.environ.get("SITE_URL", "https://whisper.chernienko.pro")

DATA_DIR = "/data"
if not os.path.exists(DATA_DIR):
    DATA_DIR = "/app/data_local"

DB_PATH = os.path.join(DATA_DIR, "users.db")
FILES_DIR = os.path.join(DATA_DIR, "files")
AUDIO_TEMP_DIR = os.path.join(DATA_DIR, "audio_cache")

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(FILES_DIR, exist_ok=True)
os.makedirs(AUDIO_TEMP_DIR, exist_ok=True)

bot = telebot.TeleBot(TG_BOT_TOKEN)
app = FastAPI(title="Whisper Pro Studio", version="2.0.0")

# --- БАЗА ДАННЫХ ---
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS login_sessions
                 (token TEXT PRIMARY KEY, user_id TEXT, created_at REAL)''')
    c.execute('''CREATE TABLE IF NOT EXISTS tasks
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  user_id TEXT, 
                  filename TEXT, 
                  status TEXT, 
                  result_path TEXT, 
                  created_at TEXT)''')
    conn.commit()
    conn.close()

init_db()

def db_update_status(task_id, status_str, result_path=None):
    try:
        conn = sqlite3.connect(DB_PATH)
        if result_path:
            conn.execute("UPDATE tasks SET status = ?, result_path = ? WHERE id = ?", (status_str, result_path, task_id))
        else:
            conn.execute("UPDATE tasks SET status = ? WHERE id = ?", (status_str, task_id))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Error updating task status: {e}")

def get_user_id_from_token(token: str):
    if not token:
        return None
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT user_id FROM login_sessions WHERE token = ?", (token,))
        row = c.fetchone()
        conn.close()
        return str(row[0]) if row and row[0] else None
    except Exception:
        return None

def unload_memory(obj=None):
    import gc
    if obj:
        del obj
    gc.collect()

def send_file_to_tg(user_id, filepath, caption):
    try:
        if os.path.exists(filepath):
            with open(filepath, "rb") as f:
                bot.send_document(user_id, f, caption=caption)
    except Exception as e:
        print(f"Error sending file to TG: {e}")

def extract_audio(input_file_path):
    """
    Мгновенно извлекает звук через ffmpeg в 16kHz mono WAV для Whisper.
    """
    base_name = os.path.splitext(os.path.basename(input_file_path))[0]
    out_wav = os.path.join(AUDIO_TEMP_DIR, f"{base_name}_{uuid.uuid4().hex[:6]}.wav")
    
    cmd = [
        "ffmpeg", "-y", "-i", input_file_path,
        "-vn",
        "-ac", "1",
        "-ar", "16000",
        "-c:a", "pcm_s16le",
        out_wav
    ]
    try:
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
        if os.path.exists(out_wav) and os.path.getsize(out_wav) > 0:
            return out_wav
    except Exception as e:
        print(f"ffmpeg extraction fallback: {e}")
    return input_file_path

# --- ОБРАБОТКА ТРАНСКРИБАЦИИ ---
def process_single_file(user_id, file_path, original_name, model_size, task_id):
    model = None
    extracted_audio = None
    try:
        from faster_whisper import WhisperModel

        db_update_status(task_id, "⏳ 1/3 Извлечение звуковой дорожки...")
        extracted_audio = extract_audio(file_path)

        db_update_status(task_id, f"⏳ 2/3 Инициализация модели ({model_size})...")
        model = WhisperModel(model_size, device="cpu", compute_type="int8", cpu_threads=4)

        db_update_status(task_id, "⏳ 3/3 Расшифровка: 0%")
        segments, info = model.transcribe(
            extracted_audio,
            language="ru",
            beam_size=1,
            vad_filter=True,
            vad_parameters=dict(min_silence_duration_ms=500),
            condition_on_previous_text=False,
            repetition_penalty=1.2,
            no_speech_threshold=0.6
        )

        duration = info.duration
        full_text = []
        last_update = 0
        for s in segments:
            t_start = time.strftime("%M:%S", time.gmtime(s.start))
            text_chunk = s.text.strip()
            if text_chunk:
                full_text.append(f"[{t_start}] — {text_chunk}")

            curr_time = time.time()
            if duration > 0 and (curr_time - last_update > 3):
                percent = min(99, int((s.end / duration) * 100))
                status_str = f"⏳ 3/3 Расшифровано {percent}% ({int(s.end)}/{int(duration)} сек)"
                db_update_status(task_id, status_str)
                last_update = curr_time

        db_update_status(task_id, "⏳ Формирование документов...")
        
        # 1. Сохраняем текстовый вариант
        txt_filename = f"Transcription_{int(time.time())}_{task_id}.txt"
        txt_path = os.path.join(FILES_DIR, txt_filename)
        
        doc = Document()
        if full_text:
            content_str = "\n".join(full_text)
            caption = f"✅ Готово ({model_size}): {original_name}"
            final_status = "✅ Готово"
        else:
            content_str = "⚠️ В аудиофайле не обнаружена речь (тишина, фоновый шум или отключенный микрофон)."
            caption = f"⚠️ Готово (речь не обнаружена): {original_name}"
            final_status = "✅ Готово (тишина)"

        with open(txt_path, "w", encoding="utf-8") as tf:
            tf.write(f"Файл: {original_name}\nМодель: {model_size}\nДата: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n" + content_str)

        # 2. Сохраняем Word .docx
        doc.add_paragraph(f"Файл: {original_name}\nМодель: {model_size}\nДата: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n" + content_str)
        res_filename = f"Transcription_{int(time.time())}_{task_id}.docx"
        res_path = os.path.join(FILES_DIR, res_filename)
        doc.save(res_path)

        db_update_status(task_id, "⏳ Отправка в Telegram...", res_path)
        send_file_to_tg(user_id, res_path, caption)
        db_update_status(task_id, final_status, res_path)

    except Exception as e:
        import traceback
        print(f"Error in process_single_file: {e}")
        traceback.print_exc()
        db_update_status(task_id, f"❌ Ошибка: {str(e)[:40]}")
    finally:
        if extracted_audio and extracted_audio != file_path and os.path.exists(extracted_audio):
            try:
                os.remove(extracted_audio)
            except Exception:
                pass
        unload_memory(model)

def process_merged_batch(user_id, file_list, model_size, task_id):
    model = None
    try:
        from faster_whisper import WhisperModel

        db_update_status(task_id, f"⏳ Инициализация модели ({model_size})...")
        model = WhisperModel(model_size, device="cpu", compute_type="int8", cpu_threads=4)
        doc = Document()
        doc.add_paragraph(f"Сводный отчет (Файлов: {len(file_list)})\nДата: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        last_update = 0
        for idx, (f_path, f_name) in enumerate(file_list):
            doc.add_page_break()
            doc.add_heading(f"Файл {idx+1}/{len(file_list)}: {f_name}", level=1)

            db_update_status(task_id, f"⏳ Обработка {idx+1}/{len(file_list)}: {f_name}...")
            extracted = extract_audio(f_path)

            segments, info = model.transcribe(
                extracted,
                language="ru",
                beam_size=1,
                vad_filter=True,
                vad_parameters=dict(min_silence_duration_ms=500),
                condition_on_previous_text=False,
                repetition_penalty=1.2,
                no_speech_threshold=0.6
            )
            duration = info.duration
            file_has_text = False
            for s in segments:
                t_start = time.strftime("%M:%S", time.gmtime(s.start))
                text_chunk = s.text.strip()
                if text_chunk:
                    doc.add_paragraph(f"[{t_start}] — {text_chunk}")
                    file_has_text = True

                curr_time = time.time()
                if duration > 0 and (curr_time - last_update > 3):
                    file_progress = s.end / duration
                    overall_progress = (idx + file_progress) / len(file_list)
                    percent = min(99, int(overall_progress * 100))
                    status_str = f"⏳ Файл {idx+1}/{len(file_list)} — {percent}% ({int(s.end)}/{int(duration)} сек)"
                    db_update_status(task_id, status_str)
                    last_update = curr_time

            if not file_has_text:
                doc.add_paragraph("⚠️ В аудиофайле не обнаружена речь (тишина или отключенный микрофон).")

            if extracted != f_path and os.path.exists(extracted):
                try:
                    os.remove(extracted)
                except Exception:
                    pass

        db_update_status(task_id, "⏳ Создание сводного отчета...")
        res_filename = f"MERGED_{int(time.time())}_{task_id}.docx"
        res_path = os.path.join(FILES_DIR, res_filename)
        doc.save(res_path)

        db_update_status(task_id, "⏳ Отправка в Telegram...", res_path)
        send_file_to_tg(user_id, res_path, "🔥 Сводный отчет готов")
        db_update_status(task_id, "✅ Пакет готов", res_path)

    except Exception as e:
        import traceback
        print(f"Error in process_merged_batch: {e}")
        traceback.print_exc()
        db_update_status(task_id, f"❌ Ошибка: {str(e)[:40]}")
    finally:
        unload_memory(model)

# --- ФОНОВАЯ ОЧЕРЕДЬ ЗАДАЧ ---
task_queue = queue.Queue()

def recover_pending_tasks():
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT id, user_id, filename, status FROM tasks WHERE status LIKE '⏳%' OR status = 'Очередь'")
        rows = c.fetchall()
        for task_id, user_id, original_name, status in rows:
            matched_file = None
            if os.path.exists(FILES_DIR):
                for fname in os.listdir(FILES_DIR):
                    if fname.endswith(original_name):
                        matched_file = os.path.join(FILES_DIR, fname)
                        break

            if matched_file and os.path.exists(matched_file):
                print(f"🔄 Восстановление задачи {task_id}: {original_name}")
                c.execute("UPDATE tasks SET status = 'Очередь' WHERE id = ?", (task_id,))
                conn.commit()
                task_queue.put({
                    "type": "single",
                    "user_id": user_id,
                    "file_path": matched_file,
                    "original_name": original_name,
                    "model_size": "medium",
                    "task_id": task_id
                })
            else:
                c.execute("UPDATE tasks SET status = '❌ Прервано' WHERE id = ?", (task_id,))
                conn.commit()
        conn.close()
    except Exception as e:
        print(f"Error in recover_pending_tasks: {e}")

def worker_loop():
    while True:
        try:
            task_info = task_queue.get()
            if task_info is None:
                break

            task_type = task_info.get("type")
            user_id = task_info.get("user_id")
            task_id = task_info.get("task_id")
            model_size = task_info.get("model_size", "medium")

            if task_type == "single":
                file_path = task_info.get("file_path")
                original_name = task_info.get("original_name")
                process_single_file(user_id, file_path, original_name, model_size, task_id)
            elif task_type == "batch":
                file_list = task_info.get("file_list")
                process_merged_batch(user_id, file_list, model_size, task_id)

            task_queue.task_done()
        except Exception as e:
            print(f"Error in background worker loop: {e}")
            time.sleep(2)

# --- TELEGRAM BOT ---
def bot_polling():
    while True:
        try:
            bot.polling(none_stop=True, interval=2, timeout=20)
        except Exception as e:
            print(f"Bot polling exception: {e}")
            time.sleep(5)

@bot.message_handler(commands=['start'])
def handle_start(message):
    try:
        args = message.text.split()
        if len(args) > 1:
            login_token = args[1].strip()
            user_id = str(message.chat.id)
            conn = sqlite3.connect(DB_PATH)
            conn.execute("INSERT OR REPLACE INTO login_sessions (token, user_id, created_at) VALUES (?, ?, ?)",
                         (login_token, user_id, time.time()))
            conn.commit()
            conn.close()
            
            callback_url = f"{SITE_URL}/login/callback?token={login_token}"
            markup = InlineKeyboardMarkup()
            markup.add(InlineKeyboardButton("🚀 Открыть личный кабинет", url=callback_url))
            bot.reply_to(
                message,
                "✅ **Вход успешно подтверждён!**\n\n"
                "Ваш браузер уже переходит в личный кабинет. Вы также можете открыть его по кнопке ниже:",
                reply_markup=markup,
                parse_mode="Markdown"
            )
        else:
            markup = InlineKeyboardMarkup()
            markup.add(InlineKeyboardButton("🌐 Открыть сайт Whisper Pro", url=SITE_URL))
            bot.reply_to(message, f"👋 Привет! Отправьте мне аудио или видео файл (до 20 МБ), либо перейдите на сайт {SITE_URL} для файлов любого размера (до 2 ГБ).", reply_markup=markup)
    except Exception as e:
        print(f"Error in handle_start: {e}")

@bot.message_handler(content_types=['audio', 'voice', 'video', 'document'])
def handle_incoming_file(message):
    try:
        user_id = str(message.chat.id)
        file_info = None
        file_name = "audio_file"

        if message.audio:
            file_info = bot.get_file(message.audio.file_id)
            file_name = message.audio.file_name or f"audio_{int(time.time())}.mp3"
        elif message.voice:
            file_info = bot.get_file(message.voice.file_id)
            file_name = f"voice_{int(time.time())}.ogg"
        elif message.video:
            file_info = bot.get_file(message.video.file_id)
            file_name = message.video.file_name or f"video_{int(time.time())}.mp4"
        elif message.document:
            file_info = bot.get_file(message.document.file_id)
            file_name = message.document.file_name or f"file_{int(time.time())}"

        if not file_info:
            bot.reply_to(message, "⚠️ Не удалось получить информацию о файле.")
            return

        bot.reply_to(message, f"📥 Файл «{file_name}» получен! Добавлен в очередь на транскрибацию...")

        downloaded_bytes = bot.download_file(file_info.file_path)
        safe_name = f"{int(time.time())}_{uuid.uuid4().hex[:4]}_{file_name}"
        saved_path = os.path.join(FILES_DIR, safe_name)

        with open(saved_path, 'wb') as new_file:
            new_file.write(downloaded_bytes)

        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("INSERT INTO tasks (user_id, filename, status, created_at) VALUES (?, ?, ?, ?)",
                       (user_id, file_name, "Очередь", datetime.now().strftime("%Y-%m-%d %H:%M")))
        task_id = cursor.lastrowid
        conn.commit()
        conn.close()

        task_queue.put({
            "type": "single",
            "user_id": user_id,
            "file_path": saved_path,
            "original_name": file_name,
            "model_size": "medium",
            "task_id": task_id
        })

    except Exception as e:
        import traceback
        print(f"Error in handle_incoming_file: {e}")
        traceback.print_exc()
        if "too big" in str(e).lower() or "file is too big" in str(e).lower():
            markup = InlineKeyboardMarkup()
            markup.add(InlineKeyboardButton("🌐 Загрузить на сайте (до 2 ГБ)", url=SITE_URL))
            bot.reply_to(
                message,
                "⚠️ **Файл превышает лимит Telegram (20 МБ)**\n\n"
                f"Загрузите файл через наш сайт {SITE_URL} — там поддерживаются любые файлы до 2 ГБ, а готовый .docx отчёт сразу придёт сюда в Telegram!",
                reply_markup=markup,
                parse_mode="Markdown"
            )
        else:
            bot.reply_to(message, f"❌ Ошибка при приеме файла: {str(e)[:50]}")

# --- FASTAPI AUTH & API ENDPOINTS ---

def get_current_user_id(request: Request):
    # 1. Cookie
    cookie_token = request.cookies.get("whisper_token")
    if cookie_token:
        uid = get_user_id_from_token(cookie_token)
        if uid:
            return uid
    # 2. Header
    auth_header = request.headers.get("Authorization")
    if auth_header and auth_header.startswith("Bearer "):
        bearer_token = auth_header.replace("Bearer ", "").strip()
        uid = get_user_id_from_token(bearer_token)
        if uid:
            return uid
    # 3. Query param token fallback
    query_token = request.query_params.get("token")
    if query_token:
        uid = get_user_id_from_token(query_token)
        if uid:
            return uid
    return None

@app.get("/login/callback")
async def login_callback(token: str):
    """Прямой редирект из Telegram: ставит cookie на 30 дней и перенаправляет в кабинет"""
    uid = get_user_id_from_token(token)
    response = RedirectResponse(url="/", status_code=302)
    if uid:
        response.set_cookie(
            key="whisper_token",
            value=token,
            max_age=30 * 86400,
            path="/",
            httponly=False,
            samesite="lax"
        )
    return response

@app.get("/api/auth/token")
async def get_auth_token():
    token = str(uuid.uuid4())
    tg_url = f"https://t.me/{TG_BOT_USERNAME}?start={token}"
    return {"token": token, "tg_url": tg_url, "bot_username": TG_BOT_USERNAME}

@app.get("/api/auth/status")
async def check_auth_status(token: str, response: Response):
    uid = get_user_id_from_token(token)
    if uid:
        response.set_cookie(
            key="whisper_token",
            value=token,
            max_age=30 * 86400,
            path="/",
            httponly=False,
            samesite="lax"
        )
        return {"authenticated": True, "user_id": uid, "token": token}
    return {"authenticated": False}

@app.post("/api/auth/logout")
async def logout(response: Response):
    response.delete_cookie("whisper_token", path="/")
    return {"status": "ok"}

@app.get("/api/tasks")
async def get_tasks(request: Request):
    uid = get_current_user_id(request)
    if not uid:
        raise HTTPException(status_code=401, detail="Не авторизован")

    conn = sqlite3.connect(DB_PATH)
    tasks = conn.execute(
        "SELECT id, created_at, filename, status, result_path FROM tasks WHERE user_id = ? ORDER BY id DESC LIMIT 30",
        (uid,)
    ).fetchall()
    conn.close()

    task_list = []
    for t in tasks:
        tid, created_at, filename, status, result_path = t
        download_docx = None
        download_txt = None
        has_text = False

        if result_path and os.path.exists(result_path):
            basename = os.path.basename(result_path)
            download_docx = f"/download/{basename}"
            txt_base = basename.replace(".docx", ".txt")
            if os.path.exists(os.path.join(FILES_DIR, txt_base)):
                download_txt = f"/download/{txt_base}"
                has_text = True

        percent = None
        match = re.search(r'(\d+)%', status)
        if match:
            percent = int(match.group(1))
        elif "✅" in status or "Готово" in status:
            percent = 100

        task_list.append({
            "id": tid,
            "created_at": created_at,
            "filename": filename,
            "status": status,
            "percent": percent,
            "download_docx": download_docx,
            "download_txt": download_txt,
            "has_text": has_text
        })

    active_task = None
    if task_list and ("⏳" in task_list[0]["status"] or "Очередь" in task_list[0]["status"]):
        active_task = task_list[0]

    return {
        "user_id": uid,
        "active_task": active_task,
        "tasks": task_list
    }

@app.get("/api/tasks/{task_id}/text")
async def get_task_text(task_id: int, request: Request):
    uid = get_current_user_id(request)
    if not uid:
        raise HTTPException(status_code=401, detail="Не авторизован")

    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT filename, result_path FROM tasks WHERE id = ? AND user_id = ?", (task_id, uid)).fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    fname, rpath = row
    if not rpath or not os.path.exists(rpath):
        return {"filename": fname, "text": "Текст еще не сформирован"}

    txt_path = rpath.replace(".docx", ".txt")
    if os.path.exists(txt_path):
        with open(txt_path, "r", encoding="utf-8") as f:
            return {"filename": fname, "text": f.read()}

    # Fallback to reading docx
    try:
        doc = Document(rpath)
        full_txt = "\n".join([p.text for p in doc.paragraphs])
        return {"filename": fname, "text": full_txt}
    except Exception as e:
        return {"filename": fname, "text": f"Ошибка чтения документа: {e}"}

@app.post("/api/tasks/create")
async def create_task(
    request: Request,
    files: list[UploadFile] = File(...),
    model_size: str = Form("medium"),
    merge_mode: bool = Form(False)
):
    uid = get_current_user_id(request)
    if not uid:
        raise HTTPException(status_code=401, detail="Не авторизован")

    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="Файлы не выбраны")

    saved_files = []
    for f in files:
        safe_name = f"{int(time.time())}_{uuid.uuid4().hex[:4]}_{os.path.basename(f.filename)}"
        saved_path = os.path.join(FILES_DIR, safe_name)
        with open(saved_path, "wb") as buffer:
            shutil.copyfileobj(f.file, buffer)
        saved_files.append((saved_path, f.filename))

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    if merge_mode and len(saved_files) > 1:
        cursor.execute("INSERT INTO tasks (user_id, filename, status, created_at) VALUES (?, ?, ?, ?)",
                       (uid, f"ПАКЕТ ({len(saved_files)} файлов)", "Очередь", datetime.now().strftime("%Y-%m-%d %H:%M")))
        task_id = cursor.lastrowid
        conn.commit()
        conn.close()

        task_queue.put({
            "type": "batch",
            "user_id": uid,
            "file_list": saved_files,
            "model_size": model_size,
            "task_id": task_id
        })
    else:
        conn.commit()
        conn.close()
        for path, name in saved_files:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            cursor.execute("INSERT INTO tasks (user_id, filename, status, created_at) VALUES (?, ?, ?, ?)",
                           (uid, name, "Очередь", datetime.now().strftime("%Y-%m-%d %H:%M")))
            task_id = cursor.lastrowid
            conn.commit()
            conn.close()

            task_queue.put({
                "type": "single",
                "user_id": uid,
                "file_path": path,
                "original_name": name,
                "model_size": model_size,
                "task_id": task_id
            })

    return {"status": "ok", "count": len(saved_files)}

@app.get("/download/{filename}")
async def download_file(filename: str):
    safe_name = os.path.basename(filename)
    file_path = os.path.join(FILES_DIR, safe_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Файл не найден")
    
    media_type = "application/octet-stream"
    if safe_name.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif safe_name.endswith(".txt"):
        media_type = "text/plain; charset=utf-8"

    return FileResponse(
        file_path,
        media_type=media_type,
        filename=safe_name
    )

@app.post("/asr")
async def api_asr(audio_file: UploadFile = File(...)):
    temp_path = os.path.join(DATA_DIR, f"n8n_{audio_file.filename}")
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(audio_file.file, buffer)

    extracted_path = None
    try:
        from faster_whisper import WhisperModel
        extracted_path = extract_audio(temp_path)
        model = WhisperModel("medium", device="cpu", compute_type="int8", cpu_threads=4)
        segments, info = model.transcribe(
            extracted_path,
            language="ru",
            beam_size=1,
            vad_filter=True,
            vad_parameters=dict(min_silence_duration_ms=500),
            condition_on_previous_text=False,
            repetition_penalty=1.2,
            no_speech_threshold=0.6
        )
        text = "".join(s.text for s in segments)
        return {"text": text}
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass
        if extracted_path and extracted_path != temp_path and os.path.exists(extracted_path):
            try:
                os.remove(extracted_path)
            except Exception:
                pass

# --- PRODUCTION-GRADE WEB APP FRONTEND (HTML5 + MODERN CSS + JS) ---
HTML_APP = """<!DOCTYPE html>
<html lang="ru">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Whisper Pro Studio — Сверхточная транскрибация речи</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&family=JetBrains+Mono:wght@400;600&display=swap" rel="stylesheet">
  <style>
    /* DESIGN_VARIANCE: Medium | MOTION_INTENSITY: Expressive | VISUAL_DENSITY: Standard */
    :root {
      --bg-main: #080c14;
      --bg-card: rgba(18, 24, 38, 0.75);
      --bg-card-hover: rgba(28, 38, 58, 0.85);
      --border-card: rgba(255, 255, 255, 0.08);
      --border-accent: rgba(59, 130, 246, 0.4);
      --primary: #2563eb;
      --primary-hover: #1d4ed8;
      --accent-cyan: #06b6d4;
      --accent-emerald: #10b981;
      --text-main: #f8fafc;
      --text-muted: #94a3b8;
      --font-sans: 'Plus Jakarta Sans', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
      --font-mono: 'JetBrains Mono', monospace;
      --radius-lg: 18px;
      --radius-md: 12px;
      --radius-sm: 8px;
      --shadow-ambient: 0 12px 36px rgba(0, 0, 0, 0.4);
    }
    * { box-sizing: border-box; margin: 0; padding: 0; }
    body {
      background: radial-gradient(circle at 50% 0%, #172554 0%, var(--bg-main) 60%);
      background-attachment: fixed;
      color: var(--text-main);
      font-family: var(--font-sans);
      min-height: 100vh;
      display: flex;
      flex-direction: column;
      overflow-x: hidden;
    }
    .container {
      max-width: 1080px;
      margin: 0 auto;
      padding: 24px 20px;
      width: 100%;
    }
    .header {
      display: flex;
      justify-content: space-between;
      align-items: center;
      padding-bottom: 20px;
      border-bottom: 1px solid var(--border-card);
      margin-bottom: 28px;
    }
    .logo-box {
      display: flex;
      align-items: center;
      gap: 14px;
      text-decoration: none;
    }
    .logo-icon {
      width: 44px;
      height: 44px;
      background: linear-gradient(135deg, var(--primary), var(--accent-cyan));
      border-radius: var(--radius-md);
      display: flex;
      align-items: center;
      justify-content: center;
      font-size: 24px;
      box-shadow: 0 4px 18px rgba(37, 99, 235, 0.4);
    }
    .logo-text h1 {
      font-size: 21px;
      font-weight: 800;
      letter-spacing: -0.5px;
      background: linear-gradient(135deg, #ffffff, #cbd5e1);
      -webkit-background-clip: text;
      -webkit-text-fill-color: transparent;
    }
    .logo-text p {
      font-size: 12px;
      color: var(--text-muted);
    }
    .status-badge {
      display: inline-flex;
      align-items: center;
      gap: 8px;
      background: rgba(15, 23, 42, 0.8);
      border: 1px solid var(--border-card);
      padding: 6px 14px;
      border-radius: 30px;
      font-size: 12px;
      color: #cbd5e1;
      font-weight: 600;
    }
    .dot {
      width: 8px;
      height: 8px;
      border-radius: 50%;
      background: var(--accent-emerald);
      box-shadow: 0 0 10px var(--accent-emerald);
    }
    .glass-card {
      background: var(--bg-card);
      backdrop-filter: blur(20px);
      -webkit-backdrop-filter: blur(20px);
      border: 1px solid var(--border-card);
      border-radius: var(--radius-lg);
      padding: 26px;
      box-shadow: var(--shadow-ambient);
      margin-bottom: 24px;
      transition: border-color 0.2s ease;
    }
    .glass-card:hover {
      border-color: rgba(255, 255, 255, 0.14);
    }

    /* AUTH SCREEN */
    .auth-card {
      max-width: 520px;
      margin: 40px auto;
      text-align: center;
      padding: 44px 30px;
    }
    .auth-icon {
      font-size: 54px;
      margin-bottom: 16px;
      display: inline-block;
      animation: floatIcon 3s ease-in-out infinite;
    }
    @keyframes floatIcon {
      0%, 100% { transform: translateY(0); }
      50% { transform: translateY(-6px); }
    }
    .auth-btn {
      display: inline-flex;
      align-items: center;
      justify-content: center;
      gap: 12px;
      background: linear-gradient(135deg, #0284c7, #0369a1);
      color: #ffffff;
      padding: 16px 36px;
      border-radius: var(--radius-md);
      text-decoration: none;
      font-weight: 700;
      font-size: 16px;
      box-shadow: 0 6px 22px rgba(2, 132, 199, 0.45);
      transition: all 0.2s ease;
      cursor: pointer;
      border: none;
      margin: 22px 0 14px 0;
      width: 100%;
    }
    .auth-btn:hover {
      transform: translateY(-2px);
      box-shadow: 0 8px 28px rgba(2, 132, 199, 0.65);
    }
    .spinner {
      display: inline-block;
      width: 18px;
      height: 18px;
      border: 2px solid rgba(255,255,255,0.25);
      border-radius: 50%;
      border-top-color: #38bdf8;
      animation: spin 0.8s linear infinite;
    }
    @keyframes spin { to { transform: rotate(360deg); } }

    /* UPLOAD DROPZONE */
    .dropzone {
      border: 2px dashed rgba(59, 130, 246, 0.4);
      background: rgba(11, 17, 33, 0.7);
      border-radius: var(--radius-md);
      padding: 40px 20px;
      text-align: center;
      cursor: pointer;
      transition: all 0.25s ease;
      position: relative;
    }
    .dropzone:hover, .dropzone.dragover {
      border-color: #38bdf8;
      background: rgba(30, 58, 138, 0.25);
      transform: scale(1.005);
    }
    .dropzone-icon {
      font-size: 46px;
      margin-bottom: 12px;
      display: block;
    }
    .file-input {
      position: absolute;
      top: 0; left: 0; width: 100%; height: 100%;
      opacity: 0;
      cursor: pointer;
    }
    .file-preview {
      margin-top: 14px;
      font-size: 14px;
      color: #38bdf8;
      font-weight: 600;
      background: rgba(56, 189, 248, 0.1);
      padding: 10px 14px;
      border-radius: var(--radius-sm);
      display: inline-block;
    }
    .controls-grid {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 18px;
      margin-top: 22px;
    }
    @media (max-width: 640px) {
      .controls-grid { grid-template-columns: 1fr; }
    }
    .form-group label {
      display: block;
      font-size: 13px;
      font-weight: 700;
      color: var(--text-muted);
      margin-bottom: 8px;
    }
    .select-input {
      width: 100%;
      background: #0b1120;
      border: 1px solid rgba(255, 255, 255, 0.12);
      border-radius: var(--radius-sm);
      color: #ffffff;
      padding: 12px 14px;
      font-size: 14px;
      font-weight: 600;
      outline: none;
    }
    .select-input:focus {
      border-color: var(--primary);
    }
    .checkbox-container {
      display: flex;
      align-items: center;
      gap: 12px;
      margin-top: 30px;
      cursor: pointer;
      font-size: 14px;
      color: #e2e8f0;
      font-weight: 600;
      user-select: none;
    }
    .btn-submit {
      width: 100%;
      background: linear-gradient(135deg, var(--primary), #1e40af);
      color: #ffffff;
      padding: 16px;
      border-radius: var(--radius-md);
      font-size: 16px;
      font-weight: 800;
      border: none;
      cursor: pointer;
      box-shadow: 0 6px 20px rgba(37, 99, 235, 0.4);
      margin-top: 24px;
      transition: all 0.2s ease;
      display: flex;
      align-items: center;
      justify-content: center;
      gap: 10px;
    }
    .btn-submit:hover:not(:disabled) {
      transform: translateY(-2px);
      box-shadow: 0 8px 26px rgba(37, 99, 235, 0.6);
    }
    .btn-submit:disabled {
      opacity: 0.6;
      cursor: not-allowed;
    }

    /* LIVE PROGRESS CARD */
    .active-card {
      border: 1px solid var(--border-accent);
      background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95));
      animation: pulseGlow 3s infinite alternate;
    }
    @keyframes pulseGlow {
      0% { box-shadow: 0 8px 24px rgba(37, 99, 235, 0.2); }
      100% { box-shadow: 0 8px 32px rgba(6, 182, 212, 0.35); }
    }
    .progress-bar-container {
      background: #141c2e;
      border: 1px solid #334155;
      border-radius: var(--radius-md);
      height: 34px;
      overflow: hidden;
      position: relative;
      margin-top: 14px;
      box-shadow: inset 0 2px 5px rgba(0,0,0,0.6);
    }
    .progress-fill {
      background: linear-gradient(90deg, #2563eb, #06b6d4, #10b981);
      height: 100%;
      transition: width 0.4s ease-in-out;
      box-shadow: 0 0 16px rgba(6, 182, 212, 0.7);
    }
    .progress-text {
      position: absolute;
      top: 0; left: 0; width: 100%; height: 100%;
      display: flex;
      align-items: center;
      justify-content: center;
      font-size: 13px;
      font-weight: 700;
      color: #ffffff;
      text-shadow: 0 1px 4px rgba(0,0,0,0.9);
      font-family: var(--font-mono);
    }
    .btn-dl-ready {
      display: inline-flex;
      align-items: center;
      gap: 8px;
      background: linear-gradient(135deg, #10b981, #059669);
      color: #ffffff;
      padding: 10px 22px;
      border-radius: var(--radius-sm);
      text-decoration: none;
      font-weight: 700;
      font-size: 14px;
      box-shadow: 0 4px 14px rgba(16, 185, 129, 0.4);
      margin-top: 14px;
    }

    /* TABS */
    .tabs {
      display: flex;
      gap: 12px;
      margin-bottom: 20px;
      border-bottom: 1px solid var(--border-card);
      padding-bottom: 8px;
    }
    .tab-btn {
      background: transparent;
      border: none;
      color: var(--text-muted);
      font-weight: 700;
      font-size: 15px;
      padding: 10px 20px;
      border-radius: var(--radius-sm);
      cursor: pointer;
      transition: all 0.2s;
    }
    .tab-btn.active {
      background: rgba(59, 130, 246, 0.15);
      color: #60a5fa;
      border-bottom: 2px solid #60a5fa;
    }

    /* TABLE */
    .table-container {
      overflow-x: auto;
      margin-top: 14px;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      font-size: 13px;
    }
    th {
      text-align: left;
      padding: 12px 14px;
      color: var(--text-muted);
      font-weight: 700;
      border-bottom: 1px solid var(--border-card);
    }
    td {
      padding: 14px;
      border-bottom: 1px solid rgba(255, 255, 255, 0.04);
      color: #e2e8f0;
    }
    .status-pill {
      display: inline-block;
      padding: 5px 12px;
      border-radius: 20px;
      font-weight: 700;
      font-size: 12px;
    }
    .status-pill.done { background: #132e1e; color: #4ade80; border: 1px solid #166534; }
    .status-pill.working { background: #2d2616; color: #fbbf24; border: 1px solid #78350f; }
    .status-pill.error { background: #331919; color: #f87171; border: 1px solid #991b1b; }
    .action-group {
      display: flex;
      align-items: center;
      gap: 8px;
    }
    .btn-table-action {
      display: inline-flex;
      align-items: center;
      gap: 4px;
      background: #1e293b;
      border: 1px solid rgba(255, 255, 255, 0.1);
      color: #cbd5e1;
      padding: 6px 12px;
      border-radius: var(--radius-sm);
      text-decoration: none;
      font-weight: 600;
      font-size: 12px;
      cursor: pointer;
      transition: all 0.2s;
    }
    .btn-table-action:hover {
      background: var(--primary);
      color: #ffffff;
      border-color: var(--primary);
    }
    .btn-table-primary {
      background: #2563eb;
      color: #ffffff;
      border: none;
    }
    .btn-table-primary:hover {
      background: #1d4ed8;
    }

    /* MODAL VIEWER */
    .modal-overlay {
      position: fixed;
      top: 0; left: 0; width: 100%; height: 100%;
      background: rgba(0, 0, 0, 0.8);
      backdrop-filter: blur(8px);
      z-index: 1000;
      display: flex;
      align-items: center;
      justify-content: center;
      padding: 20px;
    }
    .modal-card {
      background: #0f172a;
      border: 1px solid rgba(255, 255, 255, 0.15);
      border-radius: var(--radius-lg);
      max-width: 840px;
      width: 100%;
      max-height: 85vh;
      display: flex;
      flex-direction: column;
      box-shadow: 0 20px 50px rgba(0,0,0,0.7);
    }
    .modal-header {
      display: flex;
      justify-content: space-between;
      align-items: center;
      padding: 18px 24px;
      border-bottom: 1px solid var(--border-card);
    }
    .modal-body {
      padding: 24px;
      overflow-y: auto;
      font-family: var(--font-mono);
      font-size: 13px;
      line-height: 1.7;
      color: #cbd5e1;
      white-space: pre-wrap;
      user-select: text;
    }
    .modal-footer {
      display: flex;
      justify-content: flex-end;
      gap: 12px;
      padding: 16px 24px;
      border-top: 1px solid var(--border-card);
    }
    .btn-secondary {
      background: #334155;
      color: #ffffff;
      padding: 8px 18px;
      border-radius: var(--radius-sm);
      border: none;
      cursor: pointer;
      font-weight: 600;
      font-size: 13px;
    }
    .logout-btn {
      background: transparent;
      border: 1px solid rgba(255, 255, 255, 0.15);
      color: var(--text-muted);
      padding: 6px 14px;
      border-radius: var(--radius-sm);
      cursor: pointer;
      font-size: 12px;
      font-weight: 600;
      transition: all 0.2s;
    }
    .logout-btn:hover {
      background: rgba(239, 68, 68, 0.2);
      color: #f87171;
      border-color: #ef4444;
    }
    .hidden { display: none !important; }
  </style>
</head>
<body>
  <div class="container">
    <header class="header">
      <a href="/" class="logo-box">
        <div class="logo-icon">🎙</div>
        <div class="logo-text">
          <h1>Whisper Pro Studio</h1>
          <p>Сверхточная транскрибация речи на базе нейросетей</p>
        </div>
      </a>
      <div style="display:flex;align-items:center;gap:12px;">
        <div class="status-badge">
          <span class="dot"></span>
          <span>Сервер онлайн</span>
        </div>
        <button id="logoutBtn" class="logout-btn hidden" onclick="handleLogout()">🚪 Выйти</button>
      </div>
    </header>

    <!-- 1. ЭКРАН АВТОРИЗАЦИИ ЧЕРЕЗ TELEGRAM -->
    <div id="authScreen" class="glass-card auth-card">
      <span class="auth-icon">🔐</span>
      <h2 style="font-size:24px;margin-bottom:8px;font-weight:800;">Вход в Whisper Pro</h2>
      <p style="color:#94a3b8;font-size:14px;line-height:1.5;">Авторизуйтесь через ваш Telegram-аккаунт для мгновенного доступа к транскрибациям и автоматической отправки готовых файлов в чат:</p>
      
      <div>
        <a id="tgLoginBtn" href="#" target="_blank" class="auth-btn">
          ✈️ Войти через Telegram-бота (@whisper_log_bot)
        </a>
      </div>

      <div style="display:flex;align-items:center;justify-content:center;gap:10px;color:#94a3b8;font-size:13px;margin-top:14px;">
        <span class="spinner"></span>
        <span>Ожидание нажатия Start в боте... Вход выполнится <b>автоматически</b>.</span>
      </div>
    </div>

    <!-- 2. ЛИЧНЫЙ КАБИНЕТ (ПОСЛЕ ВХОДА) -->
    <div id="cabinetScreen" class="hidden">
      <!-- Активная задача (Live Card) -->
      <div id="activeTaskBox" class="glass-card active-card hidden">
        <div style="display:flex;justify-content:space-between;align-items:center;">
          <div style="font-weight:800;font-size:15px;color:#e2e8f0;">
            🔥 В процессе: <span id="activeFileName" style="color:#38bdf8;"></span>
          </div>
          <span id="activeTime" style="font-size:12px;color:#94a3b8;font-family:var(--font-mono);"></span>
        </div>
        <div class="progress-bar-container">
          <div id="progressFill" class="progress-fill" style="width:0%;"></div>
          <div id="progressText" class="progress-text">⏳ Инициализация...</div>
        </div>
        <div id="activeDlBox" style="text-align:right;"></div>
      </div>

      <!-- Вкладки -->
      <div class="tabs">
        <button id="tabNewBtn" class="tab-btn active" onclick="switchTab('new')">🚀 Новая транскрибация</button>
        <button id="tabHistoryBtn" class="tab-btn" onclick="switchTab('history')">📜 Мои расшифровки (<span id="taskCount">0</span>)</button>
      </div>

      <!-- Вкладка: Загрузка -->
      <div id="tabNew" class="glass-card">
        <div class="dropzone" id="dropzone">
          <span class="dropzone-icon">📁</span>
          <h3 style="font-size:17px;font-weight:800;margin-bottom:4px;">Перетащите аудио или видео файлы сюда</h3>
          <p style="font-size:13px;color:#94a3b8;">Поддерживаются: MP4, WEBM, MP3, OGG, WAV, M4A, MOV, MKV до 2 ГБ</p>
          <input type="file" id="fileInput" class="file-input" multiple accept="audio/*,video/*,.webm,.mp4,.ogg,.mp3,.wav,.m4a,.mov,.mkv">
          <div id="filePreview" class="file-preview hidden"></div>
        </div>

        <div class="controls-grid">
          <div class="form-group">
            <label>🧠 Модель нейросети</label>
            <select id="modelSelect" class="select-input">
              <option value="medium" selected>medium (Максимальная точность — рекомендуется)</option>
              <option value="small">small (Быстрая обработка коротких заметок)</option>
            </select>
          </div>
          <div class="form-group">
            <label class="checkbox-container">
              <input type="checkbox" id="mergeCheckbox" style="width:18px;height:18px;">
              <span>📑 Объединить файлы в один общий .docx отчет</span>
            </label>
          </div>
        </div>

        <button id="submitBtn" class="btn-submit" onclick="submitFiles()">
          <span>🚀 Запустить транскрибацию</span>
        </button>

        <div id="uploadStatusText" style="text-align:center;margin-top:14px;font-size:13px;color:#38bdf8;font-weight:700;"></div>
      </div>

      <!-- Вкладка: История -->
      <div id="tabHistory" class="glass-card hidden">
        <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:14px;">
          <h3 style="font-size:17px;font-weight:800;">📜 История расшифровок</h3>
          <button class="logout-btn" onclick="fetchTasks()">🔄 Обновить</button>
        </div>
        <div class="table-container">
          <table>
            <thead>
              <tr>
                <th>Дата</th>
                <th>Файл</th>
                <th>Статус</th>
                <th>Действия</th>
              </tr>
            </thead>
            <tbody id="historyTableBody">
              <tr><td colspan="4" style="text-align:center;color:#64748b;padding:24px;">Загрузка истории...</td></tr>
            </tbody>
          </table>
        </div>
      </div>

    </div>
  </div>

  <!-- МОДАЛЬНОЕ ОКНО ПРОСМОТРА ТЕКСТА -->
  <div id="textModal" class="modal-overlay hidden">
    <div class="modal-card">
      <div class="modal-header">
        <h3 id="modalTitle" style="font-size:16px;font-weight:700;color:#ffffff;">Текст расшифровки</h3>
        <button class="btn-secondary" onclick="closeModal()">✕</button>
      </div>
      <div id="modalContent" class="modal-body"></div>
      <div class="modal-footer">
        <button class="btn-table-action" onclick="copyModalText()">📋 Копировать текст</button>
        <button class="btn-secondary" onclick="closeModal()">Закрыть</button>
      </div>
    </div>
  </div>

  <script>
    let currentAuthToken = localStorage.getItem('whisper_session_token') || '';
    let pollAuthTimer = null;
    let pollTasksTimer = null;

    window.addEventListener('DOMContentLoaded', async () => {
      // 1. Проверяем URL параметр ?token=
      const urlParams = new URLSearchParams(window.location.search);
      const urlToken = urlParams.get('token');
      if (urlToken) {
        currentAuthToken = urlToken;
        localStorage.setItem('whisper_session_token', urlToken);
        window.history.replaceState({}, document.title, window.location.pathname);
      }

      await initAuthFlow();
      setupDropzone();
    });

    async function initAuthFlow() {
      // Попробуем запросить задачи напрямую (если есть куки)
      try {
        const res = await fetch('/api/tasks');
        if (res.ok) {
          const data = await res.json();
          renderCabinet(data);
          startTaskPolling();
          return;
        }
      } catch (e) {}

      // Если в localStorage сохранен токен, проверим его статус
      if (currentAuthToken) {
        try {
          const res = await fetch(`/api/auth/status?token=${currentAuthToken}`);
          const data = await res.json();
          if (data.authenticated) {
            const taskRes = await fetch('/api/tasks');
            if (taskRes.ok) {
              renderCabinet(await taskRes.json());
              startTaskPolling();
              return;
            }
          }
        } catch (e) {}
      }

      // Не авторизован -> показываем экран логина
      showLoginScreen();
    }

    async function showLoginScreen() {
      document.getElementById('authScreen').classList.remove('hidden');
      document.getElementById('cabinetScreen').classList.add('hidden');
      document.getElementById('logoutBtn').classList.add('hidden');

      if (pollTasksTimer) clearInterval(pollTasksTimer);

      try {
        const res = await fetch('/api/auth/token');
        const data = await res.json();
        currentAuthToken = data.token;
        document.getElementById('tgLoginBtn').href = data.tg_url;

        // Поллинг подтверждения
        if (pollAuthTimer) clearInterval(pollAuthTimer);
        pollAuthTimer = setInterval(async () => {
          try {
            const stRes = await fetch(`/api/auth/status?token=${currentAuthToken}`);
            const stData = await stRes.json();
            if (stData.authenticated) {
              clearInterval(pollAuthTimer);
              localStorage.setItem('whisper_session_token', currentAuthToken);
              const taskRes = await fetch('/api/tasks');
              if (taskRes.ok) {
                renderCabinet(await taskRes.json());
                startTaskPolling();
              }
            }
          } catch(e){}
        }, 1500);
      } catch (e) {
        console.error('Error fetching auth token:', e);
      }
    }

    function renderCabinet(data) {
      document.getElementById('authScreen').classList.add('hidden');
      document.getElementById('cabinetScreen').classList.remove('hidden');
      document.getElementById('logoutBtn').classList.remove('hidden');

      updateActiveTask(data.active_task);
      updateHistoryTable(data.tasks);
    }

    function updateActiveTask(task) {
      const box = document.getElementById('activeTaskBox');
      if (!task) {
        box.classList.add('hidden');
        return;
      }
      box.classList.remove('hidden');
      document.getElementById('activeFileName').textContent = task.filename;
      document.getElementById('activeTime').textContent = task.created_at;

      const fill = document.getElementById('progressFill');
      const txt = document.getElementById('progressText');
      txt.textContent = task.status;

      let p = task.percent;
      if (p === null || p === undefined) {
        p = task.status.includes('Очередь') ? 5 : (task.status.includes('Извлечение') ? 15 : 25);
      }
      fill.style.width = Math.min(100, Math.max(5, p)) + '%';

      const dlBox = document.getElementById('activeDlBox');
      if (task.download_docx) {
        dlBox.innerHTML = `
          <div style="display:flex;gap:10px;justify-content:flex-end;margin-top:12px;">
            <button class="btn-table-action" onclick="viewTaskText(${task.id})">👁 Просмотреть текст</button>
            <a href="${task.download_docx}" download class="btn-dl-ready">📥 Скачать DOCX</a>
          </div>
        `;
      } else {
        dlBox.innerHTML = '';
      }
    }

    function updateHistoryTable(tasks) {
      document.getElementById('taskCount').textContent = tasks ? tasks.length : 0;
      const tbody = document.getElementById('historyTableBody');
      if (!tasks || tasks.length === 0) {
        tbody.innerHTML = '<tr><td colspan="4" style="text-align:center;color:#64748b;padding:24px;">Нет созданных задач</td></tr>';
        return;
      }

      tbody.innerHTML = tasks.map(t => {
        let pillClass = 'working';
        if (t.status.includes('✅') || t.status.includes('Готово')) pillClass = 'done';
        else if (t.status.includes('❌') || t.status.includes('Ошибка')) pillClass = 'error';

        let actions = '';
        if (t.download_docx) {
          actions = `
            <div class="action-group">
              <button class="btn-table-action" onclick="viewTaskText(${t.id})">👁 Текст</button>
              <a href="${t.download_docx}" download class="btn-table-action btn-table-primary">📥 DOCX</a>
              ${t.download_txt ? `<a href="${t.download_txt}" download class="btn-table-action">📄 TXT</a>` : ''}
            </div>
          `;
        } else {
          actions = '<span style="color:#64748b;">—</span>';
        }

        return `<tr>
          <td style="color:#94a3b8;font-family:var(--font-mono);font-size:12px;">${t.created_at}</td>
          <td style="font-weight:700;color:#ffffff;">${t.filename}</td>
          <td><span class="status-pill ${pillClass}">${t.status}</span></td>
          <td>${actions}</td>
        </tr>`;
      }).join('');
    }

    async function fetchTasks() {
      try {
        const res = await fetch('/api/tasks');
        if (res.ok) {
          const data = await res.json();
          renderCabinet(data);
        }
      } catch (e){}
    }

    function startTaskPolling() {
      if (pollTasksTimer) clearInterval(pollTasksTimer);
      pollTasksTimer = setInterval(fetchTasks, 2500);
    }

    function switchTab(tab) {
      if (tab === 'new') {
        document.getElementById('tabNew').classList.remove('hidden');
        document.getElementById('tabHistory').classList.add('hidden');
        document.getElementById('tabNewBtn').classList.add('active');
        document.getElementById('tabHistoryBtn').classList.remove('active');
      } else {
        document.getElementById('tabNew').classList.add('hidden');
        document.getElementById('tabHistory').classList.remove('hidden');
        document.getElementById('tabNewBtn').classList.remove('active');
        document.getElementById('tabHistoryBtn').classList.add('active');
        fetchTasks();
      }
    }

    function setupDropzone() {
      const dropzone = document.getElementById('dropzone');
      const fileInput = document.getElementById('fileInput');
      const preview = document.getElementById('filePreview');

      ['dragenter', 'dragover'].forEach(name => {
        dropzone.addEventListener(name, (e) => { e.preventDefault(); dropzone.classList.add('dragover'); });
      });
      ['dragleave', 'drop'].forEach(name => {
        dropzone.addEventListener(name, (e) => { e.preventDefault(); dropzone.classList.remove('dragover'); });
      });

      fileInput.addEventListener('change', () => {
        if (fileInput.files.length > 0) {
          preview.classList.remove('hidden');
          preview.textContent = `Выбрано файлов: ${fileInput.files.length} шт. (` + Array.from(fileInput.files).map(f => f.name).join(', ') + ')';
        } else {
          preview.classList.add('hidden');
        }
      });
    }

    async function submitFiles() {
      const fileInput = document.getElementById('fileInput');
      if (!fileInput.files || fileInput.files.length === 0) {
        alert('Пожалуйста, выберите файлы для загрузки');
        return;
      }

      const submitBtn = document.getElementById('submitBtn');
      const statusText = document.getElementById('uploadStatusText');
      submitBtn.disabled = true;

      const formData = new FormData();
      for (let i = 0; i < fileInput.files.length; i++) {
        formData.append('files', fileInput.files[i]);
      }
      formData.append('model_size', document.getElementById('modelSelect').value);
      formData.append('merge_mode', document.getElementById('mergeCheckbox').checked);

      const xhr = new XMLHttpRequest();
      xhr.open('POST', '/api/tasks/create', true);

      xhr.upload.onprogress = (e) => {
        if (e.lengthComputable) {
          const percent = Math.round((e.loaded / e.total) * 100);
          statusText.textContent = `⏳ Загрузка на сервер: ${percent}%...`;
        }
      };

      xhr.onload = async () => {
        submitBtn.disabled = false;
        if (xhr.status === 200) {
          statusText.textContent = '🚀 Успешно отправлено в обработку! Результат появится ниже и придёт в Telegram.';
          fileInput.value = '';
          document.getElementById('filePreview').classList.add('hidden');
          await fetchTasks();
          setTimeout(() => { statusText.textContent = ''; }, 4000);
        } else {
          statusText.textContent = '❌ Ошибка при загрузке: ' + xhr.responseText;
        }
      };

      xhr.onerror = () => {
        submitBtn.disabled = false;
        statusText.textContent = '❌ Ошибка сети при отправке файлов.';
      };

      xhr.send(formData);
    }

    async function viewTaskText(taskId) {
      try {
        const res = await fetch(`/api/tasks/${taskId}/text`);
        if (res.ok) {
          const data = await res.json();
          document.getElementById('modalTitle').textContent = `Текст: ${data.filename}`;
          document.getElementById('modalContent').textContent = data.text;
          document.getElementById('textModal').classList.remove('hidden');
        }
      } catch(e) {
        alert('Ошибка при получении текста: ' + e);
      }
    }

    function closeModal() {
      document.getElementById('textModal').classList.add('hidden');
    }

    function copyModalText() {
      const txt = document.getElementById('modalContent').textContent;
      navigator.clipboard.writeText(txt).then(() => {
        alert('Текст скопирован в буфер обмена!');
      });
    }

    async function handleLogout() {
      localStorage.removeItem('whisper_session_token');
      try {
        await fetch('/api/auth/logout', { method: 'POST' });
      } catch (e) {}
      showLoginScreen();
    }
  </script>
</body>
</html>
"""

@app.get("/", response_class=HTMLResponse)
async def serve_index(request: Request):
    return HTMLResponse(content=HTML_APP)

def start_background_services():
    if getattr(start_background_services, "_started", False):
        return
    start_background_services._started = True
    print("🚀 Starting background worker loop & recovering tasks...")
    threading.Thread(target=worker_loop, daemon=True).start()
    recover_pending_tasks()
    try:
        print("Clearing active webhooks...")
        bot.remove_webhook()
    except Exception as e:
        print(f"Failed to remove webhook: {e}")
    print("🤖 Starting Telegram bot polling...")
    threading.Thread(target=bot_polling, daemon=True).start()

if __name__ == "__main__":
    start_background_services()
    uvicorn.run(app, host="0.0.0.0", port=7860)
